from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from flask_cors import CORS
from PIL import Image
from vit_keras import vit
from pydub import AudioSegment
import tensorflow as tf
import numpy as np
import transformers
import tensorflow_hub as hub
import tensorflow_text as text
from pdfminer import high_level
import librosa
import os
import docx

UPLOAD_FOLDER = "./static"
ALLOWED_EXTENSIONS = {"txt", "docx", "pdf", "jpeg", "jpg", "png", "mp3", "wav", "ogg"}

app = Flask(__name__)
cors = CORS(app)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

def convert_docx_to_txt(docx_path):
  doc = docx.Document(docx_path)
  text = ''
  for paragraph in doc.paragraphs:
    text += paragraph.text + '\n'
  return text

def convert_pdf_to_text(pdf_path):
  text = ''
  with open(pdf_path, "rb") as pdf:
    parser = high_level.PDFParser(pdf)
    document = high_level.PDFDocument(parser)
    converter = high_level.TextConverter(document)
    text = converter.convert()
  return text

def check_ext(filename):
  ext = filename.split(".")
  ext = ext[-1].lower()
  if ext in ALLOWED_EXTENSIONS:
    return True
  return False

def convert_image(filename):
  img = Image.open(f"./static/images/{filename}")
  gray = 0
  if img.mode == "RGBA" or img.mode == "P": img = img.convert("RGB")
  if img.mode == "L" or len(img.getpixel((0, 0))) == 1: gray = 1
  name = filename.split(".")
  del name[-1]
  name = ".".join(name)
  img.save(f"./static/images/{name}.jpg", quality=75, optimize=True)
  if filename != f"{name}.jpg": os.remove(f"./static/images/{filename}")
  return {"filename": f"{name}.jpg", "gray": gray}

def convert_audio(filename):
  name = filename.split(".")
  ext = name[-1].lower()
  name = ".".join(name)
  if ext == "wav": return filename
  if ext == "mp3":
    audio = AudioSegment.from_mp3(f"./static/audio/{filename}")
    audio.export(f"./static/audio/{name}.wav", format="wav")
    os.remove(f"./static/audio/{filename}")
  if ext == "ogg":
    audio = AudioSegment.from_ogg(f"./static/audio/{filename}")
    audio.export(f"./static/audio/{name}.wav", format="wav")
    os.remove(f"./static/audio/{filename}")
  return f"{name}.wav"

def preprocess_image(filename, isGray):
  img = tf.io.read_file(f"./static/images/{filename}")
  if isGray:
    img = tf.image.decode_jpeg(img, channels=1)
    img = tf.image.grayscale_to_rgb(img)
  else:
    img = tf.image.decode_jpeg(img, channels=3)
  img = tf.image.convert_image_dtype(img, tf.float32)
  img = tf.image.resize(img, size=[224, 224])
  img = np.expand_dims(img, axis=0)
  return img

def preprocess_audio(filename):
  arr = []
  waveform, sample_rate = librosa.load(f"./static/audio/{filename}", sr=None)
  mfcc = librosa.feature.mfcc(y=waveform, sr=sample_rate,
                              n_mfcc=25, n_fft=4096, hop_length=512)
  mfcc = tf.image.resize(np.expand_dims(mfcc, -1), (96, 64), method="nearest")
  mfcc = tf.expand_dims(mfcc, axis=-1)
  arr.append(mfcc)
  arr = np.array(arr)
  arr = arr.reshape(-1, 96, 64, 1)
  return arr
  
def load_and_predict_text(text):
  text_model = tf.keras.models.load_model("./models/trained_bert.h5", {'KerasLayer': hub.KerasLayer})
  text = np.expand_dims(text, axis=0)
  prediction = text_model.predict(text)
  prediction = prediction.tolist()[0][0]
  return prediction

def load_and_predict_image(filename):
  image_model = tf.keras.models.load_model("./models/NewnetV1_trained.h5")
  obj = convert_image(filename)
  img = preprocess_image(obj["filename"], obj["gray"])
  prediction = image_model.predict(img)
  prediction = prediction.tolist()[0][0]
  return prediction

def load_and_predict_audio(filename):
  audio_model = tf.keras.models.load_model("./models/vggish_model_new.h5")
  filename = convert_audio(filename)
  audio = preprocess_audio(filename)
  prediction = audio_model.predict(audio)
  prediction = prediction.tolist()[0][0]
  return prediction

@app.route("/predict", methods=["POST"])
def predict_deepfake():
  if request.method == "POST":
    f = request.files.get("file")
    if f and check_ext(f.filename):
      filename = secure_filename(f.filename)
      ext = filename.split(".")[-1].lower()
      if ext in {"txt", "docx", "pdf"}:
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], "text", filename)
        f.save(file_path)
        if ext == "docx":
          text = convert_docx_to_txt(file_path) 
        if ext == "pdf":
          text = convert_pdf_to_text(file_path) 
        if ext == "txt":
          with open(file_path, "r") as fi:
            text = fi.readlines()
        prediction = load_and_predict_text((text))
        prediction = 1 - prediction

      if filename.split(".")[-1].lower() in {"jpg", "jpeg", "png"}:
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], "images", filename)
        f.save(file_path)
        prediction = load_and_predict_image(filename)

      if filename.split(".")[-1].lower() in {"ogg", "wav", "mp3"}:
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], "audio", filename)
        f.save(file_path)
        prediction = load_and_predict_audio(filename)
    return jsonify({"prediction": prediction * 100})
